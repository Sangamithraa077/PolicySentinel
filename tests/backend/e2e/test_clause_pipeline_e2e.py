"""End-to-end verification of the full document-processing pipeline:
parsing -> normalization -> segmentation -> storage -> API.

No single HTTP endpoint chains these stages together yet -- each is its
own service (see services/document_parsing_service.py's docstring), so
this test runs parsing/normalization/segmentation directly in Python
against a realistic .docx built to exercise every stage at once:

  - a title repeated three times and page-number footers (normalization
    has noise to remove)
  - a heading -> subheading -> numbered-list-item -> bullet hierarchy,
    plus a second top-level heading whose body is a table (segmentation
    has real structure to build, in a document with *no* blank lines
    anywhere -- the tight, no-blank-line-between-markers case that's
    the trickiest part of the segmentation algorithm)

Storage and the read API are then exercised for real, over HTTP,
against the real Postgres test database (see conftest.py) -- proving
the six pipeline stages actually compose, not just that each one works
in isolation (already covered by their own unit/integration suites).
"""

from __future__ import annotations

import io
import uuid

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from backend.parsing.docx_parser import DocxDocumentParser
from backend.repositories.clause_repository import ClauseRepository
from backend.services.clause_segmentation_service import ClauseSegmentationService
from backend.services.store_segmented_clauses_service import StoreSegmentedClausesService
from backend.services.text_normalization_service import TextNormalizationService

UPLOAD_URL = "/api/v1/uploads/policies"
CLAUSES_URL = "/api/v1/clauses"
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
REPEATED_TITLE = "Acme Corp Policy Manual"


def _make_pipeline_docx() -> bytes:
    document = DocxDocument()
    for paragraph in (
        REPEATED_TITLE,
        "1. Introduction",
        "This   policy    governs   conduct.",
        "Page 1 of 3",
        REPEATED_TITLE,
        "1.1 Scope",
        "Applies to all staff.",
        "(a) First requirement",
        "(b) Second requirement",
        "- Detail one",
        "- Detail two",
        "Page 2 of 3",
        REPEATED_TITLE,
        "2. Definitions",
    ):
        document.add_paragraph(paragraph)

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Meaning"
    table.cell(1, 0).text = "Policy"
    table.cell(1, 1).text = "A governing rule"

    document.add_paragraph("Page 3 of 3")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def pipeline_result(client: TestClient, db_session: Session, seeded_company_and_user):
    """Runs the whole pipeline once (upload -> parse -> normalize ->
    segment -> store -> list via the real API); every test below
    asserts against this single run so a failure in one stage is easy
    to tell apart from a failure in another."""
    company, user = seeded_company_and_user
    docx_bytes = _make_pipeline_docx()

    upload_response = client.post(
        UPLOAD_URL,
        data={
            "company_id": str(company.id),
            "uploaded_by_user_id": str(user.id),
            "policy_title": "Pipeline Verification Policy",
            "version_number": "1",
        },
        files={"file": ("policy.docx", docx_bytes, DOCX_CONTENT_TYPE)},
    )
    assert upload_response.status_code == 201
    upload = upload_response.json()

    raw_text = DocxDocumentParser().parse(docx_bytes)
    normalized = TextNormalizationService().normalize(raw_text)
    clauses = ClauseSegmentationService().segment(normalized.text)

    StoreSegmentedClausesService(ClauseRepository(db_session)).store(
        clauses,
        policy_id=uuid.UUID(upload["policy_id"]),
        policy_version_id=uuid.UUID(upload["policy_version_id"]),
    )

    list_response = client.get(CLAUSES_URL, params={"policy_id": upload["policy_id"]})
    assert list_response.status_code == 200

    return {
        "upload": upload,
        "raw_text": raw_text,
        "normalized_text": normalized.text,
        "clauses": clauses,
        "api_items": list_response.json()["items"],
    }


# --- Document parsing ---------------------------------------------------------


@pytest.mark.text_extraction
def test_pipeline_parsing_extracts_paragraphs_and_table_from_the_docx(pipeline_result) -> None:
    raw = pipeline_result["raw_text"]

    assert "1. Introduction" in raw
    assert "1.1 Scope" in raw
    assert "Term | Meaning" in raw
    assert "Policy | A governing rule" in raw
    # This noise is normalization's job to remove, not parsing's --
    # asserting it's still here proves normalization did the work
    # rather than the parser silently stripping it.
    assert raw.count(REPEATED_TITLE) == 3
    assert "Page 1 of 3" in raw
    assert "Page 2 of 3" in raw
    assert "Page 3 of 3" in raw


# --- Text normalization ---------------------------------------------------------


@pytest.mark.text_normalization
def test_pipeline_normalization_removes_noise_and_preserves_structure(pipeline_result) -> None:
    normalized = pipeline_result["normalized_text"]

    assert normalized.count(REPEATED_TITLE) == 1
    assert "Page 1 of 3" not in normalized
    assert "Page 2 of 3" not in normalized
    assert "Page 3 of 3" not in normalized
    assert "This policy governs conduct." in normalized  # whitespace collapsed
    assert "1. Introduction" in normalized  # heading preserved
    assert "1.1 Scope" in normalized  # subheading/numbering preserved
    assert "Term | Meaning" in normalized  # table preserved as readable text


# --- Clause segmentation ---------------------------------------------------------


@pytest.mark.clause_segmentation
def test_pipeline_segmentation_builds_the_expected_hierarchy(pipeline_result) -> None:
    clauses = pipeline_result["clauses"]
    by_number = {c.clause_number: c for c in clauses if c.clause_number}

    heading = by_number["1"]
    subheading = by_number["1.1"]
    item_a = by_number["(a)"]
    item_b = by_number["(b)"]
    definitions = by_number["2"]

    assert heading.parent_id is None
    assert subheading.parent_id == heading.id
    assert item_a.parent_id == subheading.id
    assert item_b.parent_id == subheading.id
    assert definitions.parent_id is None

    bullets = [c for c in clauses if c.clause_number is None and c.text.startswith("Detail")]
    assert len(bullets) == 2
    assert all(b.parent_id == item_b.id for b in bullets)

    # Table content survives as the heading's body -- segmentation
    # doesn't try to parse table structure, just preserves the text.
    assert "Term | Meaning" in definitions.text


# --- Clause ordering ---------------------------------------------------------


@pytest.mark.clause_ordering
def test_pipeline_clauses_are_sequenced_in_document_order(pipeline_result) -> None:
    clauses = pipeline_result["clauses"]

    order_indexes = [c.order_index for c in clauses]
    assert order_indexes == list(range(len(clauses)))  # contiguous, zero-based

    numbers_in_order = [c.clause_number for c in clauses if c.clause_number]
    assert numbers_in_order == ["1", "1.1", "(a)", "(b)", "2"]

    api_order_indexes = [item["order_index"] for item in pipeline_result["api_items"]]
    assert api_order_indexes == order_indexes  # survives the storage/API round trip


# --- Database storage ---------------------------------------------------------


@pytest.mark.metadata_storage
def test_pipeline_storage_round_trips_every_clause(pipeline_result, db_session: Session) -> None:
    clauses = pipeline_result["clauses"]
    upload = pipeline_result["upload"]

    stored = ClauseRepository(db_session).list_for_policy_version(
        uuid.UUID(upload["policy_version_id"])
    )

    assert len(stored) == len(clauses)
    assert {s.id for s in stored} == {c.id for c in clauses}
    assert all(s.policy_id == uuid.UUID(upload["policy_id"]) for s in stored)
    assert all(s.policy_version_id == uuid.UUID(upload["policy_version_id"]) for s in stored)


# --- API responses ---------------------------------------------------------


@pytest.mark.api_response
def test_pipeline_api_list_matches_stored_clauses_and_hierarchy(pipeline_result) -> None:
    items = pipeline_result["api_items"]
    clauses = pipeline_result["clauses"]

    assert len(items) == len(clauses)
    assert [item["id"] for item in items] == [str(c.id) for c in clauses]

    by_number = {item["clause_number"]: item for item in items if item["clause_number"]}
    assert by_number["1.1"]["parent_clause_id"] == by_number["1"]["id"]
    assert by_number["(a)"]["parent_clause_id"] == by_number["1.1"]["id"]


@pytest.mark.api_response
def test_pipeline_api_detail_response_has_every_required_field(
    pipeline_result, client: TestClient
) -> None:
    first_item = pipeline_result["api_items"][0]

    response = client.get(f"{CLAUSES_URL}/{first_item['id']}")

    assert response.status_code == 200
    body = response.json()
    assert set(body.keys()) == {
        "id",
        "policy_id",
        "policy_version_id",
        "parent_clause_id",
        "clause_number",
        "heading",
        "text",
        "order_index",
    }


@pytest.mark.api_response
@pytest.mark.error_handling
def test_pipeline_api_returns_404_for_an_unknown_clause(client: TestClient) -> None:
    response = client.get(f"{CLAUSES_URL}/{uuid.uuid4()}")

    assert response.status_code == 404
    assert response.json()["error"]["code"] == "clause_not_found"
