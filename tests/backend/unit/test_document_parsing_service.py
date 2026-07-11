"""Unit tests for document parsing — pure logic, no I/O beyond in-memory
byte buffers, no database, no HTTP. Covers each format-specific parser,
the whitespace-cleanup helper, and the DocumentParsingService dispatcher.
"""

from __future__ import annotations

import io

import pytest
from docx import Document as DocxDocument

from backend.domain.exceptions.parsing_exceptions import DocumentParsingError
from backend.domain.exceptions.upload_exceptions import UnsupportedFileTypeError
from backend.parsing.docx_parser import DocxDocumentParser
from backend.parsing.pdf_parser import PdfDocumentParser
from backend.parsing.plain_text_parser import PlainTextDocumentParser
from backend.services.document_parsing_service import (
    DocumentParsingService,
    build_default_document_parsing_service,
)
from backend.utils.text_cleaning import clean_extracted_text


def make_minimal_pdf(*page_texts: str) -> bytes:
    """Hand-build the smallest valid multi-page PDF pypdf can extract text
    from, avoiding a dependency on a PDF-writing library just for tests."""
    objs: list[bytes] = [b"<</Type/Catalog/Pages 2 0 R>>"]
    page_count = len(page_texts)
    kids = " ".join(f"{3 + i} 0 R" for i in range(page_count))
    objs.append(f"<</Type/Pages/Kids[{kids}]/Count {page_count}>>".encode())
    content_obj_start = 3 + page_count
    font_obj = content_obj_start + page_count
    for i in range(page_count):
        objs.append(
            (
                f"<</Type/Page/Parent 2 0 R/Resources<</Font<</F1 {font_obj} 0 R>>>>"
                f"/MediaBox[0 0 200 200]/Contents {content_obj_start + i} 0 R>>"
            ).encode()
        )
    for text in page_texts:
        stream = f"BT /F1 24 Tf 10 100 Td ({text}) Tj ET".encode()
        objs.append(b"<</Length %d>>\nstream\n%s\nendstream" % (len(stream), stream))
    objs.append(b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>")

    buf = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objs, start=1):
        offsets.append(len(buf))
        buf += f"{i} 0 obj".encode() + obj + b"endobj\n"
    xref_offset = len(buf)
    buf += f"xref\n0 {len(objs) + 1}\n".encode()
    buf += b"0000000000 65535 f \n"
    for offset in offsets:
        buf += f"{offset:010d} 00000 n \n".encode()
    buf += f"trailer<</Size {len(objs) + 1}/Root 1 0 R>>\nstartxref\n{xref_offset}\n%%EOF".encode()
    return bytes(buf)


def make_docx(paragraphs: list[str]) -> bytes:
    document = DocxDocument()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_table(before: str, rows: list[list[str]], after: str) -> bytes:
    document = DocxDocument()
    document.add_paragraph(before)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for row_index, row in enumerate(rows):
        for col_index, cell_text in enumerate(row):
            table.cell(row_index, col_index).text = cell_text
    document.add_paragraph(after)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.text_extraction
def test_plain_text_parser_decodes_utf8_and_normalizes_line_endings() -> None:
    parser = PlainTextDocumentParser()

    result = parser.parse(b"Heading\r\n\r\nBody text.")

    assert result == "Heading\n\nBody text."


@pytest.mark.text_extraction
def test_plain_text_parser_falls_back_to_latin1_on_invalid_utf8() -> None:
    parser = PlainTextDocumentParser()
    content = "café".encode("latin-1")  # not valid UTF-8

    result = parser.parse(content)

    assert result == "café"


@pytest.mark.text_extraction
def test_pdf_parser_extracts_text_and_joins_pages_with_blank_line() -> None:
    parser = PdfDocumentParser()
    pdf_bytes = make_minimal_pdf("Page One", "Page Two")

    result = parser.parse(pdf_bytes)

    assert "Page One" in result
    assert "Page Two" in result
    assert result.index("Page One") < result.index("Page Two")


@pytest.mark.text_extraction
@pytest.mark.error_handling
def test_pdf_parser_raises_document_parsing_error_for_corrupt_content() -> None:
    parser = PdfDocumentParser()

    with pytest.raises(DocumentParsingError):
        parser.parse(b"not a real pdf")


@pytest.mark.text_extraction
def test_docx_parser_extracts_paragraphs_in_order() -> None:
    parser = DocxDocumentParser()
    docx_bytes = make_docx(["Section 1", "", "Body paragraph."])

    result = parser.parse(docx_bytes)

    assert result == "Section 1\n\nBody paragraph."


@pytest.mark.text_extraction
def test_docx_parser_renders_tables_as_pipe_delimited_rows_in_document_order() -> None:
    parser = DocxDocumentParser()
    docx_bytes = make_docx_with_table(
        "Before the table",
        [["Name", "Role"], ["Alice", "Engineer"]],
        "After the table",
    )

    result = parser.parse(docx_bytes)

    assert result == (
        "Before the table\nName | Role\nAlice | Engineer\nAfter the table"
    )


@pytest.mark.text_extraction
@pytest.mark.error_handling
def test_docx_parser_raises_document_parsing_error_for_corrupt_content() -> None:
    parser = DocxDocumentParser()

    with pytest.raises(DocumentParsingError):
        parser.parse(b"not a real docx")


@pytest.mark.text_extraction
def test_clean_extracted_text_collapses_internal_whitespace() -> None:
    assert clean_extracted_text("Too   many    spaces") == "Too many spaces"


@pytest.mark.text_extraction
def test_clean_extracted_text_strips_line_edges() -> None:
    assert clean_extracted_text("  leading and trailing  \n  next line  ") == (
        "leading and trailing\nnext line"
    )


@pytest.mark.text_extraction
def test_clean_extracted_text_collapses_excess_blank_lines() -> None:
    assert clean_extracted_text("Para one\n\n\n\n\nPara two") == "Para one\n\nPara two"


@pytest.mark.text_extraction
def test_clean_extracted_text_strips_document_edges() -> None:
    assert clean_extracted_text("\n\n  Hello  \n\n") == "Hello"


@pytest.mark.text_extraction
def test_service_dispatches_to_registered_parser_by_extension() -> None:
    service = DocumentParsingService({".txt": PlainTextDocumentParser()})

    result = service.parse(b"Some   policy   text.", extension=".txt")

    assert result.text == "Some policy text."
    assert result.extension == ".txt"
    assert result.character_count == len(result.text)


@pytest.mark.text_extraction
def test_service_extension_lookup_is_case_insensitive() -> None:
    service = DocumentParsingService({".txt": PlainTextDocumentParser()})

    result = service.parse(b"content", extension=".TXT")

    assert result.extension == ".txt"


@pytest.mark.text_extraction
@pytest.mark.error_handling
def test_service_raises_unsupported_file_type_for_unregistered_extension() -> None:
    service = DocumentParsingService({".txt": PlainTextDocumentParser()})

    with pytest.raises(UnsupportedFileTypeError):
        service.parse(b"content", extension=".exe")


@pytest.mark.text_extraction
def test_default_service_supports_all_four_document_types() -> None:
    service = build_default_document_parsing_service()

    assert service.parse(b"plain text", extension=".txt").text == "plain text"
    assert service.parse(b"# Heading\nBody", extension=".md").text == "# Heading\nBody"
    assert "Hello" in service.parse(make_minimal_pdf("Hello"), extension=".pdf").text
    assert service.parse(make_docx(["Hello"]), extension=".docx").text == "Hello"
